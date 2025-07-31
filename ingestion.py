import os
from typing import List, Union
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader
from langchain_chroma import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings

# New imports for file handling
import PyPDF2
import docx
import pandas as pd

load_dotenv()

api_key = os.environ.get("GEMINI_API_KEY")

embeddings = GoogleGenerativeAIEmbeddings(
    model="models/text-embedding-004",
    google_api_key=api_key,
)

def load_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    text = ""
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def load_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_csv(file_path: str) -> str:
    """Extract text from a CSV file (as concatenated rows)."""
    df = pd.read_csv(file_path)
    return df.to_string(index=False)

def load_txt(file_path: str) -> str:
    """Read text from a TXT file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def load_file(file_path: str) -> str:
    """Load text from a file based on its extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return load_pdf(file_path)
    elif ext == '.docx':
        return load_docx(file_path)
    elif ext == '.csv':
        return load_csv(file_path)
    elif ext == '.txt':
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def ingest_files(file_paths: List[str]) -> List[str]:
    """Ingest and split text from a list of files."""
    texts = [load_file(fp) for fp in file_paths]
    return texts

def process_documents(texts: List[str]):
    """Split and embed documents, returning splits."""
    text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=250, chunk_overlap=0
    )
    docs_splits = []
    for text in texts:
        docs_splits.extend(text_splitter.split_text(text))
    return docs_splits

def add_documents_to_vectorstore(docs: list):
    """
    Add new documents (as text chunks) to the Chroma vectorstore.
    """
    from langchain_core.documents import Document
    # Convert text chunks to Document objects if needed
    doc_objs = [Document(page_content=chunk) for chunk in docs]
    vectorstore = Chroma(
        collection_name="rag-chroma",
        persist_directory=".chroma",
        embedding_function=embeddings,
    )
    vectorstore.add_documents(doc_objs)

# Existing web loader logic (for reference)
urls = [
    "https://lilianweng.github.io/posts/2023-06-23-agent/",
    "https://lilianweng.github.io/posts/2023-03-15-prompt-engineering/",
    "https://lilianweng.github.io/posts/2023-10-25-adv-attack-llm/",
]

docs =  [WebBaseLoader(url).load() for url in urls]
docs_list = [item for sublist in docs for item in sublist]

text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
    chunk_size=250, chunk_overlap=0
)

docs_splits = text_splitter.split_documents(docs_list)

print(docs_splits)

# vectorstore = Chroma.from_documents(
#     documents=docs_splits,
#     collection_name="rag-chroma",
#     embedding=embeddings,
#     persist_directory=".chroma",
# )

retriever = Chroma(
    collection_name="rag-chroma",
    persist_directory=".chroma",
    embedding_function=embeddings,
).as_retriever()

if __name__ == "__main__":
    # Example usage for local files
    file_paths = ["gemini paper.pdf"]
    try:
        texts = ingest_files(file_paths)
        splits = process_documents(texts)
        add_documents_to_vectorstore(splits)
        print(f"Processed {len(splits)} document chunks.")
    except Exception as e:
        print(f"Error: {e}")

    # Example usage for retriever
    query = "What are the main differences between the prompt engineering techniques discussed in the blog posts?"
    res =  retriever.invoke(query)
    print(res)
